from docx import Document
from docx.shared import Inches

import os

document = Document()

document.add_heading('Pandat仿真报告', 0)

#p = document.add_paragraph('A plain paragraph having some ')
#p.add_run('bold').bold = True
#p.add_run(' and some ')
#p.add_run('italic.').italic = True

#document.add_heading('Heading, level 1', level=1)

#document.add_paragraph('Intense quote', style='Intense Quote')

document.add_heading('输入数据',level = 2)

#document.add_paragraph(
#    'first item in ordered list', style='List Number'
#)

recordset=[]
#selectedElements=["Cu"]
elementsFile = open("element.txt","r")
elementList = elementsFile.readlines()
elementLists = []
for fields in elementList:
    elementLists.append(fields);

contentFile1 = open("content1.txt","r")
contentList1 = contentFile1.readlines()
contentLists1 = []
for fields in contentList1:
    contentLists1.append(fields);

n=0;
for element in elementLists:
        print("element:")
        print(element)
        recordset.append({
        "ingredient" : element,
        "content": contentLists1[n]
    })
        n=n+1


#recordset = [
#    {
#        "ingredient" : "Cu",
#        "content": 98
 #   },
  #  {
   #     "ingredient" : "Fe",
    #    "content": 2
    #}
#]
#print(recordset[0]['id'])
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '成分'
hdr_cells[1].text = '含量'
for item in recordset:
	row_cells = table.add_row().cells
	row_cells[0].text = str(item['ingredient'])
	row_cells[1].text = str(item['content'])

document.add_heading('输出数据',level = 2)

fileadd = open("TTT_1.dat","r")
contentadd = fileadd.read()
document.add_paragraph(contentadd)
fileadd.close()

graog_1 = document.add_picture('WeChat Screenshot_20201231153701.png', width=Inches(2))

document.save('Pandat仿真报告.docx')
 



